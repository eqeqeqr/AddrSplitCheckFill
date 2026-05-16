from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from pypdf import PdfReader


@dataclass
class ParsedDocumentChunk:
    source_id: str
    filename: str
    file_type: str
    location: str
    text: str
    sequence: int


MAX_CHUNK_CHARS = 2000
SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".txt", ".csv", ".tsv", ".docx", ".pdf"}


def _clean_text(value: Any) -> str:
    text = "" if value is None or pd.isna(value) else str(value)
    return " ".join(text.replace("\r", "\n").split())


def _chunk_text(source_id: str, filename: str, file_type: str, location: str, text: str, start_sequence: int = 1) -> list[ParsedDocumentChunk]:
    cleaned = _clean_text(text)
    if not cleaned:
        return []

    chunks: list[ParsedDocumentChunk] = []
    for index in range(0, len(cleaned), MAX_CHUNK_CHARS):
        chunks.append(
            ParsedDocumentChunk(
                source_id=source_id,
                filename=filename,
                file_type=file_type,
                location=location if len(cleaned) <= MAX_CHUNK_CHARS else f"{location} 片段 {len(chunks) + 1}",
                text=cleaned[index : index + MAX_CHUNK_CHARS],
                sequence=start_sequence + len(chunks),
            )
        )
    return chunks


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件编码不支持，请转换为 UTF-8 或 GBK")


def _parse_text(path: Path, source_id: str) -> list[ParsedDocumentChunk]:
    text = _read_text_file(path)
    return _chunk_text(source_id, path.name, path.suffix.lower().lstrip("."), "全文", text)


def _parse_table(path: Path, source_id: str) -> list[ParsedDocumentChunk]:
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(path, sheet_name=None).items()
    else:
        sep = "\t" if suffix == ".tsv" else ","
        sheets = [("Sheet1", pd.read_csv(path, sep=sep))]

    chunks: list[ParsedDocumentChunk] = []
    sequence = 1
    for sheet_name, df in sheets:
        df = df.fillna("")
        for start in range(0, len(df), 50):
            part = df.iloc[start : start + 50]
            text = part.to_csv(index=False)
            location = f"{sheet_name} 第 {start + 1}-{start + len(part)} 行"
            next_chunks = _chunk_text(source_id, path.name, suffix.lstrip("."), location, text, sequence)
            chunks.extend(next_chunks)
            sequence += len(next_chunks)
    return chunks


def _parse_word(path: Path, source_id: str) -> list[ParsedDocumentChunk]:
    document = Document(path)
    chunks: list[ParsedDocumentChunk] = []
    sequence = 1
    for index, paragraph in enumerate(document.paragraphs, start=1):
        next_chunks = _chunk_text(source_id, path.name, "docx", f"第 {index} 段", paragraph.text, sequence)
        chunks.extend(next_chunks)
        sequence += len(next_chunks)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text for cell in row.cells))
        next_chunks = _chunk_text(source_id, path.name, "docx", f"表格 {table_index}", "\n".join(rows), sequence)
        chunks.extend(next_chunks)
        sequence += len(next_chunks)
    return chunks


def _parse_pdf(path: Path, source_id: str) -> list[ParsedDocumentChunk]:
    reader = PdfReader(path)
    chunks: list[ParsedDocumentChunk] = []
    sequence = 1
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        next_chunks = _chunk_text(source_id, path.name, "pdf", f"第 {index} 页", text, sequence)
        chunks.extend(next_chunks)
        sequence += len(next_chunks)
    return chunks


def parse_document(path: Path, source_id: str) -> list[ParsedDocumentChunk]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("资料格式不支持")
    if suffix in {".xlsx", ".xls", ".csv", ".tsv"}:
        return _parse_table(path, source_id)
    if suffix == ".txt":
        return _parse_text(path, source_id)
    if suffix == ".docx":
        return _parse_word(path, source_id)
    if suffix == ".pdf":
        return _parse_pdf(path, source_id)
    raise ValueError("资料格式不支持")
